from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

INDIGO_PURPLE = "4B0082"
LIGHT_GREY = "EDEDED"


# ---------------- HELPERS ----------------

def add_horizontal_line(paragraph, color=INDIGO_PURPLE):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_background(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_cell_padding(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, value in [("top", top), ("bottom", bottom),
                        ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_two_column_table(doc, data):
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(4.8)

    for row_idx, (label, value) in enumerate(data):
        left_cell = table.rows[row_idx].cells[0]
        right_cell = table.rows[row_idx].cells[1]

        left_run = left_cell.paragraphs[0].add_run(label)
        left_run.bold = True
        left_run.font.size = Pt(10)
        left_run.font.color.rgb = RGBColor.from_string(INDIGO_PURPLE)
        set_cell_background(left_cell, LIGHT_GREY)
        set_cell_padding(left_cell)

        right_run = right_cell.paragraphs[0].add_run(value)
        right_run.font.size = Pt(10)
        set_cell_padding(right_cell)




# ---------------- FRONT PAGE ----------------

def add_front_page(doc, meta):
    """
    meta = {
        "dut_name": "",
        "dut_version": "",
        "start_time": "",
        "end_time": "",
        "final_result": "PASS/FAIL",
        "itsar_id": "",
        "itsar_version": ""
    }
    """


    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        f"TSTL Evaluation Test Report for {os.getenv("DUT_NAME")}"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(INDIGO_PURPLE)
    add_horizontal_line(title)

    doc.add_paragraph()
    doc.add_paragraph()

    # -------- Test Details --------
    h = doc.add_paragraph("Test Details")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = RGBColor.from_string(INDIGO_PURPLE)
    add_horizontal_line(h)

    test_details = [
        ("DUT", meta["dut_name"]),
        ("Version", meta["dut_version"]),
        ("Test Category", "Common Security Requirements"),
        ("Test Case ID", "1.6.1_TC_CRYPTOGRAPHIC_BASED_SECURE_COMMUNICATION"),
        ("Type", "Auto Generated Validation Report"),
        ("Start Date", meta["start_time"]),
        ("End Date", meta["end_time"]),
        ("Requirement Test Result", meta["final_result"]),
    ]

    add_two_column_table(doc, test_details)
    doc.add_paragraph()

    # -------- DUT Information --------
    h = doc.add_paragraph("DUT Information")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = RGBColor.from_string(INDIGO_PURPLE)
    add_horizontal_line(h)

    dut_info = [
        ("DUT Details", meta["dut_name"]),
        ("DUT Software Version", meta["dut_version"]),
        ("Digest Hash of OS", meta["os_hash"]),
        ("Digest Hash of Configuration", meta["config_hash"]),
        ("Applicable ITSAR", meta["itsar_id"]),
        ("ITSAR Version No", meta["itsar_version"]),
    ]

    add_two_column_table(doc, dut_info)

    # 🔴 Important: page break after front page
    doc.add_page_break()