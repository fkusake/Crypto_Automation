from dotenv import load_dotenv
load_dotenv()
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_UNDERLINE
import os


from SSH_TEST_CASES.SSH_TC1 import run_cipher_detection
from SSH_TEST_CASES.SSH_TC2 import run_ssh_verification
from NMAP.NMAP_SCAN import run_nmap_scan
from SSH_TEST_CASES.SSH_TC3 import run_ssh_weak_cipher_test
from SSH_TEST_CASES.SSH_TC4 import run_ssh_none_cipher_test
from DUT_INFO.DUT_INFO_PAGE import add_front_page
from DUT_INFO.GET_DUT_INFO import get_dut_info
from HTTPS_TEST_CASES.HTTPS_TC1 import run_httpsCipher_detection
from HTTPS_TEST_CASES.HTTPS_TC2 import run_tls_verification
from HTTPS_TEST_CASES.HTTPS_TC3 import run_https_weak_cipher_test
from HTTPS_TEST_CASES.HTTPS_TC4 import run_https_NULL_test
from OEM_TEST_CASES.OEM_TC1 import  run_OEM_test
from SNMP_TEST_CASES.SNMP_TC1 import run_snmp_tc1
from SNMP_TEST_CASES.SNMP_TC2 import run_snmp_tc2


REPORTS = "REPORT"
os.makedirs(REPORTS, exist_ok=True)


# ---------------- FORMAT HELPERS ----------------
def add_itsar_subheading(doc, text, level):
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(75, 0, 130)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    return para


def add_test_case(doc, title, description):
    p = doc.add_paragraph()
    
    run1 = p.add_run(f"• {title} : ")
    run1.bold = True

    run2 = p.add_run(description)


def add_itsar_heading(doc, text, level):
    para = doc.add_paragraph()

    # Font size based on level (matches your old logic)
    font_size = Pt(16 if level == 1 else 14)

    run = para.add_run(text)
    run.bold = True
    run.font.size = font_size
    run.font.color.rgb = RGBColor(75, 0, 130)  # same blue you used

    # Spacing (tight like ITSAR image)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(2)

    # Add bottom border (horizontal line)
    p = para._p
    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')        # line thickness
    bottom.set(qn('w:space'), '2')      # gap between text and line
    bottom.set(qn('w:color'), '4B0082') # same heading blue

    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def add_bold_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    return p


def style_table_header(cell, bg_color="1F4E79"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg_color)
    tc_pr.append(shd)

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


def prevent_table_row_split(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)


def normal_screenshot_evidence_block(doc, title, image_path):
    TABLE_WIDTH = Inches(7.8)
    IMAGE_WIDTH = Inches(6.2)

    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # 🔒 PREVENT TABLE FROM BREAKING ACROSS PAGES
    prevent_table_row_split(table)

    # Lock table width
    table.columns[0].width = TABLE_WIDTH
    for row in table.rows:
        row.cells[0].width = TABLE_WIDTH

    # ---------- COMMON CELL STYLE ----------
    for row in table.rows:
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F3ECFA')
        tcPr.append(shd)

        cell.top_margin = Inches(0.2)
        cell.bottom_margin = Inches(0.2)
        cell.left_margin = Inches(0.3)
        cell.right_margin = Inches(0.3)

    # ---------- TITLE ----------
    title_cell = table.cell(0, 0)
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True

    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 0, 130)

    # ---------- IMAGE ----------
    img_cell = table.cell(1, 0)
    p_img = img_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.keep_together = True
    p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

    # ---------- PURPLE BORDER ----------
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), '4B0082')
        tblBorders.append(border)

    tblPr.append(tblBorders)

    return table

def add_screenshot_evidence_block(doc, title, image_path, overview_text=None):
    TABLE_WIDTH = Inches(7.8)
    IMAGE_WIDTH = Inches(6.2)

    table = doc.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    prevent_table_row_split(table)

    table.columns[0].width = TABLE_WIDTH
    for row in table.rows:
        row.cells[0].width = TABLE_WIDTH

    # ---------- BACKGROUNDS ----------
    for i, row in enumerate(table.rows):
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')

        if i == 2:
            shd.set(qn('w:fill'), 'F3ECFA')  # White for overview
        else:
            shd.set(qn('w:fill'), 'F3ECFA')  # Purple

        tcPr.append(shd)

        cell.top_margin = Inches(0.2)
        cell.bottom_margin = Inches(0.2)
        cell.left_margin = Inches(0.3)
        cell.right_margin = Inches(0.3)

    # ---------- TITLE ----------
    title_cell = table.cell(0, 0)
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 0, 130)

    # ---------- IMAGE ----------
    img_cell = table.cell(1, 0)
    p_img = img_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

    # ---------- OVERVIEW SECTION ----------
    if overview_text:
        overview_cell = table.cell(2, 0)

        # 🔹 Overview Title (CENTER + UNDERLINE)
        p_head = overview_cell.paragraphs[0]
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head.paragraph_format.space_after = Pt(2)   # 🔥 reduce gap

        run = p_head.add_run("Overview")
        run.bold = True
        run.font.size = Pt(11) 
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.underline = True   # ✅ proper underline

# 🔹 Overview text (CENTERED)
        p_text = overview_cell.add_paragraph()
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER   # 🔥 CENTER TEXT
        p_text.paragraph_format.space_before = Pt(2)   # 🔥 reduce gap

        run = p_text.add_run(overview_text)
        run.font.size = Pt(10)

    # ---------- BORDER ----------
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), '4B0082')
        tblBorders.append(border)

    tblPr.append(tblBorders)

    return table

def add_grey_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    p_elm = p._p
    p_pr = p_elm.get_or_add_pPr()

    p_borders = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')        # thin line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFBFBF')  # grey

    p_borders.append(bottom)
    p_pr.append(p_borders)

    return p
# -----------------------------------------------


# ---------------- UTILITY ----------------
def normalize_list(items):
    if not items:
        return ["None"]
    cleaned = [i.strip() for i in items if i.strip()]
    return cleaned if cleaned else ["None"]
# ----------------------------------------

# ---------------- FRONT PAGE ----------------



# ---------------- WORD REPORT ----------------
def create_word_report(nmap_data,OEM_data,cipher_data,ssh_data,weak_cipher_result,none_cipher_result,https_cipher_data,https_data,https_weak_cipher_result, tc8_result,dut_info,snmp_tc1_data,snmp_tc2_data,ssh_applicable,https_applicable,snmp_applicable):
    doc = Document()

    section_counter = 1
    secondsection_counter = 1
    thirdsection_counter = 1

    start_time = datetime.now()

    # ---------------- FRONT PAGE ----------------
    tc1 = cipher_data.get("result", "FAIL")
    tc2 = ssh_data.get("final_result", "FAIL")

    final_result = "PASS" if tc1 == "PASS" and tc2 == "PASS" else "FAIL"

    add_front_page(
        doc,
        meta={
            "dut_name": dut_info["dut_name"],
            "dut_version": dut_info["dut_version"],
            "os_hash": dut_info["os_hash"],
            "config_hash": dut_info["config_hash"],
            "start_time": start_time.strftime("%Y-%m-%d %H:%M:%S"),
            "end_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "final_result": final_result,
            "itsar_id": "ITSAR111092401",
            "itsar_version": "1.0",
        }
    )

    # ================= SECTIONS 1–3 =================
    add_itsar_heading(doc, f"{section_counter}. ITSAR Section No. & Name: Section 1.6 Data Protection", 2)
    section_counter += 1

    add_itsar_heading(doc, f"{section_counter}. Security Requirement No. & Name: 1.6.1 Cryptographic Based Secure Communication", 2)
    section_counter += 1
    

    add_itsar_heading(doc, f"{section_counter}. Requirement Description", 2)
    section_counter += 1

    doc.add_paragraph(
    "The communication security dimension ensures that information flows only between "
    "authorized endpoints, preventing diversion or interception during transmission. "
    "The data is protected against well-known attacks such as sniffing, disclosure, and reconnaissance."
    )

    doc.add_paragraph(
    "Secure communication mechanisms between the CPE and connected entities shall use "
    "industry-standard protocols such as IPsec, VPN, SSH, TLS/SSL, etc., along with "
    "NIST-specified cryptographic algorithms and appropriate key sizes, including SHA, "
    "Diffie-Hellman, and AES."
    )

    doc.add_paragraph(
    "Additionally, if APIs are supported, the communication between the API server and client "
    "must be strictly protected using the secure cryptographic controls prescribed in Table 1 "
    "of the latest document, 'Cryptographic Controls for Indian Telecom Security Assurance "
    "Requirements (ITSAR)'."
    )

    # ================= DUT CONFIRMATION =================
    add_itsar_heading(doc, f"{section_counter}. DUT Confirmation", 2)

    doc.add_paragraph(
    "This section is currently under development and will be updated upon completion of the "
    "DUT confirmation activities. Relevant validation details and supporting evidence will be "
    "incorporated in subsequent revisions of the report. "
    )





    # ================= DUT CONFIGURATION =================
    add_itsar_heading(doc, f"{section_counter}. DUT Configuration", 2)

    h = add_itsar_subheading(
    doc,
    f"{section_counter}.{secondsection_counter} OEM Declaration",
    2
    )
    secondsection_counter += 1
    keep_with_next(h)

    doc.add_paragraph(
    "Note: The DUT operates in standalone mode, meaning all OAM-related configurations and "
    "management functions are performed locally on the device itself. No centralized controller is used for "
    "configuration or management."
    )

    doc.add_paragraph(
    "1) The DUT supports usage of SSHv2 and HTTPS as OAM Protocols. The same has been mentioned "
    "in the Declaration shared by OEM. Console access is not available on DUT."
    )

    # ---------------- OEM DATA ----------------
    oem_data = OEM_data["oem_protocol_data"]

# Table rows = protocols + header
    table = doc.add_table(rows=len(oem_data) + 1, cols=3)
    table.style = "Table Grid"

# Header names
    headers = ["Protocol", "Supported", "Details"]

# -------- APPLY HEADER STYLING --------
    for i, header in enumerate(headers):
     cell = table.cell(0, i)

     cell.text = ""

     p = cell.paragraphs[0]
     run = p.add_run(header)

     run.bold = True
     run.font.color.rgb = RGBColor(255, 255, 255)  # White text

     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Purple background
     style_table_header(cell, bg_color="4B0082")

    # Padding
     cell.top_margin = Inches(0.15)
     cell.bottom_margin = Inches(0.15)
     cell.left_margin = Inches(0.15)
     cell.right_margin = Inches(0.15)

# -------- FILL DATA ROWS --------
    for i, (protocol, values) in enumerate(oem_data.items(), start=1):

     table.cell(i, 0).text = protocol
     table.cell(i, 1).text = values.get("supported", "Not Found")
     table.cell(i, 2).text = values.get("details", "Not Found")

# Paragraph spacing
    p = doc.add_paragraph()
    keep_with_next(p)



    image_path = os.path.join(os.path.dirname(__file__),
                          "DIAGRAMS",
                          "OEM.png")

    p = doc.add_paragraph()
    keep_with_next(p)

    p.add_run().add_picture(image_path, width=Inches(6.5))


    doc.add_paragraph(
    "2) The DUT utilizes IEEE 802.11 for wireless connectivity and DHCP for dynamic IP address allocation to "
    "connected clients."
    )

    doc.add_paragraph(
    "3) The DUT operates in standalone mode and does not rely on external control plane protocols or centralized controllers. "
    "Authentication and access control mechanisms are handled locally within the device (e.g hostpad), "
    "without the use of external RADIUS or controller-based authentication systems."
    )

    doc.add_paragraph(
    "4) The DUT does not utilize controller-based API communication. "
    "Device configuration and management are performed locally via secure shell (SSH) and web-based "
    "interface (LuCI), without reliance on centralized API-driven orchestration."
    )


    


    # ================= Protocal Verfication =================

    h = add_itsar_subheading(
    doc,
    f"{section_counter}.{secondsection_counter} Verification of Supported Communication Protocols",
    2
    )
    secondsection_counter += 1
    keep_with_next(h)

    doc.add_paragraph(
    "Nmap scanning was performed to identify the communication protocols and "
    "services supported by the DUT."
    )

# ---------------- TCP SCAN ----------------

    add_bold_paragraph(doc, "Execution Command (TCP Services):")
    doc.add_paragraph(nmap_data["user_input_tcp_ports"])

    add_bold_paragraph(doc, "Executed Command Output (TCP Services):")
    doc.add_paragraph(
    nmap_data["terminal_output_tcp_ports"] or "No output"
    )

    if nmap_data.get("tcp_screenshot") and os.path.exists(nmap_data["tcp_screenshot"]):
     normal_screenshot_evidence_block(
        doc,
        "DUT Configuration : Nmap TCP Scan Screenshot",
        nmap_data["tcp_screenshot"]
    )
     
    doc.add_paragraph()

# ---------------- UDP SCAN ----------------

    add_bold_paragraph(doc, "Execution Command (SNMP UDP Scan):")
    doc.add_paragraph(nmap_data["user_input_udp_ports"])

    add_bold_paragraph(doc, "Executed Command Output (SNMP UDP Scan):")
    doc.add_paragraph(
    nmap_data["terminal_output_udp_ports"] or "No output"
    )

    if nmap_data.get("udp_screenshot") and os.path.exists(nmap_data["udp_screenshot"]):
     normal_screenshot_evidence_block(
        doc,
        "DUT Configuration : Nmap UDP Scan Screenshot",
        nmap_data["udp_screenshot"]
    )


# ---------------- NEXT SECTION ----------------

    add_itsar_subheading(
    doc,
    f"{section_counter}.{secondsection_counter} Observation",
    2
    )

    section_counter += 1
    secondsection_counter = 1


    supported_services = []
    unsupported_services = []


# -------- Check services --------
 
    if nmap_data.get("SSH"):
     supported_services.append("SSH (Port 22)")
    else:
     unsupported_services.append("SSH (Port 22)")

    if nmap_data.get("HTTP"):
     supported_services.append("HTTP (Port 80)")
    else:
     unsupported_services.append("HTTP (Port 80)")

    if nmap_data.get("HTTPS"):
     supported_services.append("HTTPS (Port 443)")
    else:
     unsupported_services.append("HTTPS (Port 443)")

    if nmap_data.get("SNMP"):
     supported_services.append("SNMP (Port 161)")
    else:
     unsupported_services.append("SNMP (Port 161)")


    supported_text = ", ".join(supported_services)
    unsupported_text = ", ".join(unsupported_services)


    doc.add_paragraph(
    f"Based on the Nmap scan results, the Device Under Test (DUT) supports the "
    f"following communication services: {supported_text}. "
    f"The following services were not detected or were found to be closed on the DUT: "
    f"{unsupported_text}."
    )

        # ================= PRE-CONDITIONS =================
    add_itsar_heading(doc, f"{section_counter}. Preconditions", 2)
    section_counter += 1
    doc.add_paragraph(
        "• The OEM Documentation contains information about supported OAM protocols and control "
        "plane routing protocols."
    )


    image_path = os.path.join(os.path.dirname(__file__),
                          "DIAGRAMS",
                          "OEM.png")

    p = doc.add_paragraph()
    keep_with_next(p)

    p.add_run().add_picture(image_path, width=Inches(6.5))



    doc.add_paragraph(
        "• The tester has remote access to the DUT."
    )
    doc.add_paragraph(
        "• The IP address of the DUT is allocated manually by Tester."
    )

    # ================= TEST OBJECTIVE =================
    add_itsar_heading(doc, f"{section_counter}. Test Objective", 2)
    section_counter += 1
    doc.add_paragraph(
        "• To verify the secure communication between the DUT and connected entities."
    )

    # ================= TEST SCENARIO =================
    add_itsar_heading(doc, f"{section_counter}. Test Plan", 2)
    doc.add_paragraph(
        "• The DUT operates in standalone mode and supports management protocols such as SSH and HTTPS"
        "(over TLS). SNMP and IPsec are not enabled/considered in the current test scope."
    )

    doc.add_paragraph(
        "• The DUT supports SSH and HTTPS for device management, and all management interfaces are directly "
        "accessible on the DUT."
    )

    doc.add_paragraph(
        "• Secure communication between the DUT and its peers (e.g., Syslog, RADIUS, if configured) is handled "
        "locally on the device using supported cryptographic mechanisms. No centralized controller is involved."
    )

    doc.add_paragraph(
        "• The following test cases are applicable to the DUT and are executed directly on the device as part of "
        "the automated validation framework."
    )


    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Number of Test Scenarios", 2)
    secondsection_counter += 1

    add_test_case(
    doc,
    "TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS",
    "This test case verifies that the DUT supports only strong cryptographic algorithms "
    "as defined in Crypto Control ITSAR. It ensures that secure configurations are enforced on the DUT."
    )

    add_test_case(
    doc,
    "TC2_SSH_SECURE_COMMUNICATION",
    "Test to verify SSHv2 communicates securely using secure cryptographic control "
    "as provided in Crypto Control ITSAR."
    )

    add_test_case(
    doc,
    "TC3_SSH_WEAK_CIPHER_REJECTION",
    "Verify that SSHv2 communication on the DUT fails when attempting to establish a session using weak "
    "or deprecated cryptographic controls. The DUT shall reject negotiation of such weak ciphers."
    )

    add_test_case(
    doc,
    "TC4_SSH_NO_ENCRYPTION_REJECTION",
    "Verify that SSHv2 communication on the DUT fails when attempting to establish a session without encryption "
    "(e.g., none cipher). The DUT shall enforce encryption and must not allow unencrypted SSH communication."
    )

    add_test_case(
    doc,
    "TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS :",
    "This test case verifies that the DUT supports only strong TLS cipher suites as defined in Crypto Control ITSAR."
    )

    add_test_case(
    doc,
    "TC6_HTTPS_SECURE_COMMUNICATION : ",
    "Verify that HTTPS communication on the DUT is successfully established using strong TLS cipher suites "
    "as defined in Crypto Control ITSAR."
    )

    add_test_case(
    doc,
    "TC7_HTTPS_WEAK_CIPHER_REJECTION : ",
    "Verify that HTTPS communication on the DUT fails when attempting to establish a connection using weak "
    "or non-compliant TLS cipher suites. The DUT shall reject such insecure cipher suites."
    )

    add_test_case(
    doc,
    "TC8_HTTPS_NO_ENCRYPTION_REJECTION : ",
    "Verify that HTTPS communication on the DUT fails when attempting to establish a connection using null "
    "or non-encrypted cipher suites. The DUT shall enforce encryption."
    )




    image_path = os.path.join(os.path.dirname(__file__),
                          "DIAGRAMS",
                          "Untitled design.png")

    h = add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Bed Diagram", 2)
    secondsection_counter += 1
    keep_with_next(h)

    p = doc.add_paragraph()
    keep_with_next(p)

    p.add_run().add_picture(image_path, width=Inches(6.5))

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Tools Required", 2)
    section_counter += 1
    secondsection_counter = 1
    doc.add_paragraph("• Linux-based Tester System")
    doc.add_paragraph("• OpenSSH")
    doc.add_paragraph("• OpenSSL")
    doc.add_paragraph("• Nmap")
    doc.add_paragraph("• Wireshark / Tshark")
    doc.add_paragraph("• Python Automation Scripts")


    # ================= EXPECTED RESULT =================
    add_itsar_heading(doc, f"{section_counter}. Expected Results for Pass :", 2)
    section_counter += 1
    add_test_case(
    doc,
    "TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS",
    "The DUT shall support only strong cryptographic algorithms for SSHv2 communication as defined in Crypto Control ITSAR. "
    "No weak or non-compliant algorithms should be observed in the supported cipher list."
    )

    add_test_case(
    doc,
    "TC2_SSH_SECURE_COMMUNICATION",
    "The DUT should successfully establish SSHv2 communication with the Client (Tester PC) using approved secure cryptographic controls. "
    "The communication shall be encrypted using strong ciphers as defined in Crypto Control ITSAR."
    )

    add_test_case(
    doc,
    "TC3_SSH_WEAK_CIPHER_REJECTION",
    "The communication attempt shall be terminated when the Client (Tester PC) attempts to establish SSHv2 communication using weak or deprecated cryptographic algorithms. "
    "The DUT must reject such cipher negotiation."
    )

    add_test_case(
    doc,
    "TC4_SSH_NO_ENCRYPTION_REJECTION",
    "The communication attempt shall be terminated when the Client (Tester PC) attempts to establish SSHv2 communication without encryption (e.g., none cipher). "
    "The DUT must enforce encryption and disallow unencrypted communication."
    )

    add_test_case(
    doc,
    "TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS",
    "The DUT shall support only strong TLS cipher suites as defined in Crypto Control ITSAR. "
    "No weak or non-compliant cipher suites should be present in the supported list."
    )

    add_test_case(
    doc,
    "TC6_HTTPS_SECURE_COMMUNICATION",
    "The DUT should successfully establish HTTPS communication with the Client (Tester PC) using approved secure TLS cipher suites. "
    "The communication shall be encrypted and comply with Crypto Control ITSAR requirements."
    )

    add_test_case(
    doc,
    "TC7_HTTPS_WEAK_CIPHER_REJECTION",
    "The communication attempt shall be terminated when the Client (Tester PC) attempts to establish HTTPS communication using weak or non-compliant TLS cipher suites. "
    "The DUT must reject such cipher negotiation."
    )

    add_test_case(
    doc,
    "TC8_HTTPS_NO_ENCRYPTION_REJECTION",
    "The communication attempt shall be terminated when the Client (Tester PC) attempts to establish HTTPS communication using null or non-encrypted cipher suites. "
    "The DUT must enforce encryption and disallow unprotected communication."
    )



    add_itsar_heading(doc, f"{section_counter}. Expected Format of Evidence :", 2)
    section_counter += 1
    doc.add_paragraph(
       "Screenshots of respective tools (Wireshark or Terminal) used for the testing are attached during the "
       "test is performed."
    )



# SSH Section
    if nmap_data.get("SSH"):
     add_ssh_section(doc, cipher_data, ssh_data, weak_cipher_result,none_cipher_result,section_counter,secondsection_counter,thirdsection_counter)
     section_counter += 1
    else:
     add_itsar_heading(doc, f"{section_counter}. SSH Service Validation", 2)
     section_counter += 1
     doc.add_paragraph(
        "During the service discovery phase using Nmap scan, the SSH service "
        "(port 22) was not detected on the Device Under Test (DUT). "
        "Therefore, SSH-based secure communication validation tests were "
        "not executed as the DUT does not expose an SSH management interface."
     )

# HTTPS Section
    if nmap_data.get("HTTPS"):
     add_https_section(doc, https_cipher_data, https_data,https_weak_cipher_result, tc8_result,section_counter,secondsection_counter,thirdsection_counter)
     section_counter += 1
    else:
     add_itsar_heading(doc, f"{section_counter}. HTTPS Service Validation", 2)
     section_counter += 1
     doc.add_paragraph(
        "During the service discovery phase using Nmap scan, the HTTPS service "
        "(port 443) was not detected on the Device Under Test (DUT). "
        "Therefore, HTTPS-based secure communication validation tests were "
        "not executed."
    )
     
    # SNMP Section
    if nmap_data.get("SNMP"):
     add_snmp_section(doc,snmp_tc1_data,snmp_tc2_data,section_counter,secondsection_counter)
     section_counter += 1
    else:
     add_itsar_heading(doc, f"{section_counter}. SNMP Service Validation", 2)
     section_counter += 1
     doc.add_paragraph(
        "During the service discovery phase using Nmap scan, the SNMP service "
        "(port 161) was not detected on the Device Under Test (DUT). "
        "Therefore, SNMP-based secure communication validation tests were "
        "not executed."
    )
     
    # ---------- Heading ----------
    h = add_itsar_heading(doc, f"{section_counter}. Test Case Result", 2)
    section_counter += 1
    keep_with_next(h)

# ---------------- Compute Results ----------------

# -------- SSH TCs --------
    if ssh_applicable:
     tc1 = cipher_data.get("result", "NA")
     tc2 = ssh_data.get("final_result", "NA")
     tc4 = none_cipher_result.get("result", "NA")

    # TC3 (Weak Algorithm Negotiation)
     tc3 = "PASS"
    for r in weak_cipher_result.get("results", []):
        if r.get("negotiated") is True:
            tc3 = "FAIL"
            break
    else:
     tc1 = tc2 = tc3 = tc4 = "NA"

# -------- HTTPS TCs --------
    if https_applicable:
     tc5 = https_cipher_data.get("result", "NA")
     tc6 = https_data.get("final_result", "NA")
     tc8 = tc8_result.get("final_result", "NA")

    # TC7 (Weak Cipher Negotiation)
    tc7 = "PASS"
    for r in https_weak_cipher_result.get("results", []):
        if r.get("negotiated") is True:
            tc7 = "FAIL"
            break
    else:
     tc5 = tc6 = tc7 = tc8 = "NA"

    # -------- SNMP TCs --------
    if snmp_applicable:
        tc9 = snmp_tc1_data.get("final_result", "NA")
        tc10 = snmp_tc2_data.get("final_result", "NA")
    else:
        tc9 = tc10 = "NA"


# ---------------- Helper Function ----------------
    def get_remark(result, pass_msg, fail_msg, na_msg):
     if result == "PASS":
        return pass_msg
     elif result == "FAIL":
        return fail_msg
     else:
        return na_msg


# ---------------- Table ----------------
    rt = doc.add_table(rows=11, cols=4)
    rt.style = "Table Grid"

    headers = ["SL. No", "TEST CASE NAME", "PASS/FAIL", "Remarks"]

# ---------- Header ----------
    for i, header in enumerate(headers):
     cell = rt.cell(0, i)
     cell.text = ""

     p = cell.paragraphs[0]
     run = p.add_run(header)

     run.bold = True
     run.font.color.rgb = RGBColor(255, 255, 255)

     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

     style_table_header(cell, bg_color="4B0082")


# ---------------- Data ----------------
    data = [

    ("1", "TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS", tc1,
    get_remark(
     tc1,
     "Only secure SSH cryptographic algorithms are supported.",
     "Weak or unsupported SSH algorithms detected.",
     "SSH service not available — test not applicable."
    )),

    ("2", "TC2_SSH_SECURE_COMMUNICATION", tc2,
    get_remark(
     tc2,
     "SSH communication established securely using approved cryptographic controls.",
     "SSH communication uses weak or non-compliant cryptographic parameters.",
     "SSH service not available — test not applicable."
    )),

    ("3", "TC3_SSH_WEAK_CIPHER_REJECTION", tc3,
    get_remark(
     tc3,
     "DUT rejects weak SSH algorithms during negotiation.",
     "Weak SSH algorithm successfully negotiated (security issue).",
     "SSH service not available — test not applicable."
    )),

    ("4", "TC4_SSH_NO_ENCRYPTION_REJECTION", tc4,
    get_remark(
     tc4,
     "DUT rejects unencrypted SSH communication (none cipher).",
     "Unencrypted SSH communication allowed (critical vulnerability).",
     "SSH service not available — test not applicable."
    )),

    ("5", "TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS", tc5,
    get_remark(
     tc5,
     "Only secure TLS cryptographic algorithms are supported.",
     "Weak or unsupported TLS algorithms detected.",
     "HTTPS service not available — test not applicable."
    )),

    ("6", "TC6_HTTPS_SECURE_COMMUNICATION", tc6,
    get_remark(
     tc6,
     "HTTPS communication established securely using approved TLS protocol and cipher.",
     "HTTPS communication uses weak or non-compliant TLS configuration.",
     "HTTPS service not available — test not applicable."
    )),

    ("7", "TC7_HTTPS_WEAK_CIPHER_REJECTION", tc7,
    get_remark(
     tc7,
     "DUT rejects weak TLS cipher suites during negotiation.",
     "Weak TLS cipher suite successfully negotiated (security issue).",
     "HTTPS service not available — test not applicable."
    )),

    ("8", "TC8_HTTPS_NO_ENCRYPTION_REJECTION", tc8,
    get_remark(
     tc8,
     "DUT enforces encryption and rejects NULL/unencrypted TLS communication.",
     "Unencrypted TLS communication allowed (critical vulnerability).",
     "HTTPS service not available — test not applicable."
    )),

    ("9", "TC1_SNMP_VERSION_CHECK", tc9,
    get_remark(tc9,
        "Only SNMPv3 is enabled; insecure SNMPv1/v2c are disabled.",
        "SNMPv1/v2c are enabled, allowing insecure communication.",
        "SNMP service not available — test not applicable."
    )),

    ("10", "TC2_SNMP_SECURE_COMMUNICATION", tc10,
    get_remark(tc10,
        "SNMPv3 communication is secure (authPriv enforced; weak modes rejected).",
        "Weak SNMPv3 modes (authNoPriv/noAuthNoPriv) are allowed or secure mode not enforced.",
        "SNMP service not available — test not applicable."
    )),
    ]


# ---------- Fill Table ----------
    for i, row_data in enumerate(data, start=1):
     for j, value in enumerate(row_data):
        cell = rt.cell(i, j)
        cell.text = str(value)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 3 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 🔥 NA Styling (optional but recommended)
        if value == "NA":
            run = p.runs[0]
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)


# ---------- Padding ----------
    for row in rt.rows[1:]:
     for cell in row.cells:
        cell.top_margin = Inches(0.12)
        cell.bottom_margin = Inches(0.12)
        cell.left_margin = Inches(0.12)
        cell.right_margin = Inches(0.12)

    prevent_table_row_split(rt)


    # ---------- Heading ----------
    h = add_itsar_heading(doc, f"{section_counter}. Test Conclusion", 2)
    section_counter += 1
    keep_with_next(h)

    # ---------- Compute Results ----------

    # -------- SSH TC3 --------
    tc3 = "PASS"
    for r in weak_cipher_result.get("results", []):
        if r.get("negotiated") is True:
            tc3 = "FAIL"
            break

    # -------- HTTPS TC7 --------
    tc7 = "PASS"
    for r in https_weak_cipher_result.get("results", []):
        if r.get("negotiated") is True:
            tc7 = "FAIL"
            break

    # -------- SSH --------
    if ssh_applicable:
        tc1 = cipher_result.get("result", "FAIL")
        tc2 = ssh_result.get("final_result", "FAIL")
        tc4 = none_cipher_result.get("result", "FAIL")
    else:
        tc1 = tc2 = tc3 = tc4 = "NA"

    # -------- HTTPS --------
    if https_applicable:
        tc5 = https_cipher_data.get("result", "FAIL")
        tc6 = https_data.get("final_result", "FAIL")
        tc8 = tc8_result.get("final_result", "FAIL")
    else:
        tc5 = tc6 = tc7 = tc8 = "NA"

    # -------- SNMP --------
    if snmp_applicable:
        tc9 = snmp_tc1_data.get("final_result", "FAIL")
        tc10 = snmp_tc2_data.get("final_result", "FAIL")
    else:
        tc9 = tc10 = "NA"


    # ---------- Helper ----------
    def add_conclusion(doc, result, pass_msg, fail_msg, na_msg):
        p = doc.add_paragraph(style='List Bullet')

        if result == "PASS":
            p.add_run(pass_msg)
        elif result == "FAIL":
            p.add_run(fail_msg)
        else:
            p.add_run(na_msg)


    # ---------- Conclusions ----------

    # TC1
    add_conclusion(doc, tc1,
        "The DUT supports only strong SSH cryptographic algorithms and does not include weak or deprecated algorithms.",
        "The DUT supports weak or non-compliant SSH cryptographic algorithms, which violates ITSAR requirements.",
        "SSH service is not available on the DUT, hence this test case is not applicable."
    )

    # TC2
    add_conclusion(doc, tc2,
        "The DUT successfully establishes secure SSH communication using approved cryptographic mechanisms.",
        "The DUT fails to establish secure SSH communication or uses weak cryptographic parameters.",
        "SSH service is not available on the DUT, hence this test case is not applicable."
    )

    # TC3
    add_conclusion(doc, tc3,
        "The DUT rejects weak SSH algorithms during negotiation, ensuring secure communication.",
        "The DUT allows weak SSH algorithms to be negotiated, which is a security vulnerability.",
        "SSH service is not available on the DUT, hence this test case is not applicable."
    )

    # TC4
    add_conclusion(doc, tc4,
        "The DUT enforces encryption and rejects SSH connections using 'none' cipher.",
        "The DUT allows unencrypted SSH communication using 'none' cipher, which is a critical vulnerability.",
        "SSH service is not available on the DUT, hence this test case is not applicable."
    )

    # TC5
    add_conclusion(doc, tc5,
        "The DUT supports only secure TLS cryptographic algorithms for HTTPS communication.",
        "The DUT supports weak or non-compliant TLS cryptographic algorithms.",
        "HTTPS service is not available on the DUT, hence this test case is not applicable."
    )

    # TC6
    add_conclusion(doc, tc6,
        "The DUT establishes secure HTTPS communication using approved TLS protocol and cipher suites.",
        "The DUT uses weak or non-compliant TLS protocol or cipher suites during HTTPS communication.",
        "HTTPS service is not available on the DUT, hence this test case is not applicable."
    )

    # TC7
    add_conclusion(doc, tc7,
        "The DUT rejects weak TLS cipher suites during negotiation, ensuring strong encryption.",
        "The DUT allows weak TLS cipher suites to be negotiated, which is a security issue.",
        "HTTPS service is not available on the DUT, hence this test case is not applicable."
    )

    # TC8
    add_conclusion(doc, tc8,
        "The DUT enforces encryption and rejects NULL or unencrypted TLS communication.",
        "The DUT allows NULL or unencrypted TLS communication, which is a critical vulnerability.",
        "HTTPS service is not available on the DUT, hence this test case is not applicable."
    )

    # TC9 (SNMP Version Check)
    add_conclusion(doc, tc9,
        "The DUT allows only SNMPv3 communication and blocks insecure SNMPv1/v2c protocols, ensuring secure management communication.",
        "The DUT allows SNMPv1 and/or SNMPv2c communication, which exposes management data over insecure channels and violates ITSAR requirements.",
        "SNMP service is not available on the DUT, hence this test case is not applicable."
    )

    # TC10 (SNMP Secure Communication)
    add_conclusion(doc, tc10,
        "The DUT enforces secure SNMPv3 communication using authPriv mode and rejects weaker modes such as authNoPriv and noAuthNoPriv.",
        "The DUT allows SNMPv3 communication using weaker security modes (e.g., authNoPriv), which do not provide encryption and violate ITSAR security requirements.",
        "SNMP service is not available on the DUT, hence this test case is not applicable."
    )



    name = (
    "ITSAR_2.6.1_Cryptographic_Based_Secure_Communication_Report_"
    f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    report_path = os.path.join("REPORT", name)

    doc.save(report_path)
    return name


def add_ssh_section(doc, cipher_data, ssh_data, weak_cipher_result,none_cipher_result,section_counter,secondsection_counter,thirdsection_counter):

    # ================= TEST EXECUTION =================
    add_itsar_heading(doc, f"{section_counter}. Test Execution For SSH", 2)
    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 1", 2)
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph("DUT should support secure ciphers")

    add_bold_paragraph(doc, "c) Execution Steps:")
    p = doc.add_paragraph()

    p.add_run(
    "• The tester should run the command "
    )

    bold_run = p.add_run(
    "nmap --script ssh2-enum-algos <ip address>"
    )
    bold_run.bold = True

    p.add_run(
    " and check for any unsupported cipher."
    )
    doc.add_paragraph(
        "• Validate Results "
    )
    doc.add_paragraph(
        "• Capture evidence screenshots "
    )



    # ================= TC-01 =================

    add_bold_paragraph(doc, "Execution Command:")
    doc.add_paragraph(cipher_data["user_input"])

    add_bold_paragraph(doc, "Executed Command Output:")
    doc.add_paragraph(cipher_data["terminal_output"] or "No output")

    details = cipher_data.get("details", {})

    h = add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter}.{thirdsection_counter} DUT-Supported Encryption Algorithms", 2)
    thirdsection_counter += 1
    keep_with_next(h)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Strong Encryption"
    t.cell(0, 1).text = "Weak Encryption"
    row = t.add_row().cells
    row[0].text = "\n".join(normalize_list(details.get("encryption", {}).get("strong", [])))
    row[1].text = "\n".join(normalize_list(details.get("encryption", {}).get("weak", [])))
    prevent_table_row_split(t)

    h = add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter}.{thirdsection_counter} DUT-Supported MAC Algorithms", 2)
    thirdsection_counter += 1
    keep_with_next(h)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Strong MAC"
    t.cell(0, 1).text = "Weak MAC"
    row = t.add_row().cells
    row[0].text = "\n".join(normalize_list(details.get("mac", {}).get("strong", [])))
    row[1].text = "\n".join(normalize_list(details.get("mac", {}).get("weak", [])))
    prevent_table_row_split(t)

    h = add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter}.{thirdsection_counter} DUT-Supported Key Exchange Algorithms", 2)
    thirdsection_counter += 1
    keep_with_next(h)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Strong KEX"
    t.cell(0, 1).text = "Weak KEX"
    row = t.add_row().cells
    row[0].text = "\n".join(normalize_list(details.get("kex", {}).get("strong", [])))
    row[1].text = "\n".join(normalize_list(details.get("kex", {}).get("weak", [])))
    prevent_table_row_split(t)


    h = add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter}.{thirdsection_counter} DUT-Supported Host Key Algorithms", 2)
    thirdsection_counter = 1
    secondsection_counter += 1
    keep_with_next(h)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Strong Host Key"
    t.cell(0, 1).text = "Weak Host Key"
    row = t.add_row().cells
    row[0].text = "\n".join(normalize_list(details.get("host_key", {}).get("strong", [])))
    row[1].text = "\n".join(normalize_list(details.get("host_key", {}).get("weak", [])))
    prevent_table_row_split(t)
    p = doc.add_paragraph()
    keep_with_next(p)

    doc.add_paragraph() 

    # ✅ TC-01 Screenshot RESTORED
    if cipher_data.get("screenshot") and os.path.exists(cipher_data["screenshot"]):
     normal_screenshot_evidence_block(
        doc,
        "Test Case 1 : Executed cmd output: sudo nmap --script ssh2-enum-algos",
        cipher_data["screenshot"]
    )
     
    add_bold_paragraph(doc, "d) Test Observations:")
    weak_enc = ", ".join(cipher_data["details"]["encryption"]["weak"])
    weak_mac = ", ".join(cipher_data["details"]["mac"]["weak"])
    weak_kex = ", ".join(cipher_data["details"]["kex"]["weak"])
    weak_host = ", ".join(cipher_data["details"]["host_key"]["weak"])

    weak_algos = ", ".join(filter(None, [weak_enc, weak_mac, weak_kex, weak_host]))

    if weak_algos:
     observation = (
        f"The SSH algorithm enumeration results indicate the presence of weak or deprecated cryptographic algorithm(s): {weak_algos}. "
        f"As per ITSAR cryptographic requirements, the use of such weak algorithms is not permitted. "
        f"The presence of these algorithms indicates non-compliance with security standards. "
        f"Hence, the DUT is marked as FAIL for this test case."
    )
    else:
     observation = (
        "The SSH algorithm enumeration results indicate that no weak or deprecated cryptographic algorithms are supported by the DUT. "
        "All supported algorithms comply with ITSAR cryptographic requirements. "
        "Hence, the DUT is marked as PASS for this test case."
    )
    doc.add_paragraph(observation)

    add_bold_paragraph(doc, "d) Evidence Provided :")
    doc.add_paragraph("Screenshots and Packet captures are Provided during Testing.")


    # ================= TC-02 =================
    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 2", 2)
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC2_SSH_SECURE_COMMUNICATION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph("DUT and tester device should support traffic protection through SSH only through provided cryptographic methods")

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph(
        "• The tester should attempt to ssh into the DUT and capture the SSH traffic using wireshark "
    )
    doc.add_paragraph(
        "• Analyze the ssh handshake in the p-cap file."
    )

    add_bold_paragraph(doc, "Execution Command:")
    doc.add_paragraph(ssh_data["user_input"])


    h=add_itsar_subheading(doc,f"{section_counter}.{secondsection_counter}.{thirdsection_counter}.Security Ciphers Utilized During Data Transmission ( vv Fetch)", 2)
    secondsection_counter += 1
    keep_with_next(h)
    enc = doc.add_table(rows=4, cols=2)
    enc.style = "Table Grid"

    enc.cell(0, 0).text = "Protocol"
    enc.cell(0, 1).text = ssh_data["crypto_details"].get("protocol", "Not Found")

    enc.cell(1, 0).text = "Encryption Algorithm"
    enc.cell(1, 1).text = ssh_data["crypto_details"].get("cipher", "Not Found")

    enc.cell(2, 0).text = "Key Exchange Algorithm"
    enc.cell(2, 1).text = ssh_data["crypto_details"].get("kex", "Not Found")

    enc.cell(3, 0).text = "Host Key Algorithm"
    enc.cell(3, 1).text = ssh_data["crypto_details"].get("host_key", "Not Found")

    p = doc.add_paragraph()
    keep_with_next(p)


    if ssh_data.get("screenshots"):

     kex_algo = ssh_data["crypto_details"].get("kex", "Unknown")
     cipher_algo = ssh_data["crypto_details"].get("cipher", "Unknown")

    kex_status = ssh_data["nist_validation"].get("kex", "FAIL")
    enc_status = ssh_data["nist_validation"].get("encryption", "FAIL")

    kex_label = "Secure" if kex_status == "PASS" else "Insecure"
    enc_label = "Secure" if enc_status == "PASS" else "Insecure"

    titles = [
        "Test Case 2 : SSH CLI Packet Capture showing SSH Handshake Traffic",
        f"Test Case 2 : SSH Handshake showing {kex_label} Key Exchange Algorithm : {kex_algo}",
        f"Test Case 2 : SSH Data Encryption using {enc_label} Cipher : {cipher_algo}",
    ]

    heading = [
       "Overview of CL capture by T-shark",
       "Overview of Key Exchange Capture by Wireshark",
       "Overview of Encryption Capture by Wireshark"
    ]

    overviews = [

        "The above screenshot shows the command line packet capture using t-shark while performing SSH communication between the tester system and the DUT. The captured packets include the SSH handshake packets exchanged during the establishment of the SSH session.",

        f"The above screenshot shows the SSH handshake process where the Key Exchange algorithm {kex_algo} is used to securely establish the cryptographic session between the SSH client and server. According to the cryptographic ITSAR guidelines, this key exchange algorithm is classified as {kex_label}.",

        f"The above screenshot shows the encrypted SSH communication packets after the successful handshake process. The Encryption algorithm {cipher_algo} is used to encrypt the SSH data packets exchanged between the client and the DUT. According to the cryptographic ITSAR guidelines, this cipher is classified as {enc_label}."

    ]

    for idx, (img, title, overview) in enumerate(zip(
        ssh_data["screenshots"], titles, overviews)):

     if os.path.exists(img):

        add_screenshot_evidence_block(
            doc,
            title,
            img,
            overview  # 👈 inside box now
        )

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)
    

    add_bold_paragraph(doc, "d) Test Observations:")

    protocol = ssh_data["crypto_details"].get("protocol", "Unknown")
    cipher = ssh_data["crypto_details"].get("cipher", "Unknown")
    kex = ssh_data["crypto_details"].get("kex", "Unknown")
    host = ssh_data["crypto_details"].get("host_key", "Unknown")

    enc_status = ssh_data["nist_validation"].get("encryption", "FAIL")
    kex_status = ssh_data["nist_validation"].get("kex", "FAIL")

    if enc_status == "PASS" and kex_status == "PASS":
     observation = (
        f"The SSH communication between the tester and the DUT was established using protocol {protocol} "
        f"with encryption algorithm {cipher}, key exchange algorithm {kex}, and host key algorithm {host}. "
        f"All the negotiated cryptographic algorithms are compliant with ITSAR requirements. "
        f"This confirms that secure cryptographic controls are enforced during data transmission. "
        f"Hence, the DUT is marked as PASS for this test case."
    )
    else:
     observation = (
        f"The SSH communication between the tester and the DUT utilized protocol {protocol} with encryption algorithm {cipher} "
        f"and key exchange algorithm {kex}. However, one or more negotiated cryptographic algorithms are not compliant with ITSAR requirements. "
        f"This indicates that secure cryptographic controls are not fully enforced during data transmission. "
        f"Hence, the DUT is marked as FAIL for this test case."
    )

    doc.add_paragraph(observation)


    add_bold_paragraph(doc, "d) Evidence Provided :")
    doc.add_paragraph("Screenshots and Packet captures are Provided during Testing.")

# ================= WEAK CIPHER TEST (NEGATIVE) =================

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 3", 2)
    secondsection_counter += 1
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC3_SSH_WEAK_CIPHER_REJECTION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph("Verify that SSHv2 communication on the DUT fails when attempting to establish a session using weak or deprecated cryptographic controls"
    "The DUT shall reject negotiation of such weak ciphers.")

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph(
    "• The tester shall initiate an SSH connection attempt to the DUT using specified weak or deprecated cryptographic algorithms"
    )
    doc.add_paragraph(
    "• The SSH client shall be executed in verbose mode to capture detailed negotiation logs during the handshake process."
    )
    doc.add_paragraph(
    "• The verbose output shall be analyzed to determine whether the weak cryptographic algorithm is accepted or rejected by the DUT."
    )



    if weak_cipher_result.get("screenshots"):

     results = weak_cipher_result.get("results", [])
     screenshots = weak_cipher_result.get("screenshots", [])

    for idx, (res, img) in enumerate(zip(results, screenshots)):

        cmd = res.get("command", "N/A")
        terminal_output = res.get("terminal_output", "").strip()
        type_of = res.get("type", "")
        algorithm_of = res.get("algorithm", "")

        # -------- Execution Command --------
        add_bold_paragraph(doc, f"Execution Command for weak {type_of}:{algorithm_of}")
        doc.add_paragraph(cmd)

        # -------- Executed Command Output --------
        add_bold_paragraph(doc, "Executed Command Output:")

        if terminal_output:
            doc.add_paragraph(terminal_output)
        else:
            doc.add_paragraph("No relevant output captured (algorithm likely rejected during negotiation)")

        algo = res.get("algorithm", "Unknown")
        algo_type = res.get("type", "Unknown")
        negotiated = res.get("negotiated", False)

        status_label = "Insecure"
        negotiation_text = "successfully negotiated" if negotiated else "rejected by the DUT"

        # -------- Title --------
        if algo_type == "cipher":
            title = f"SSH Weak Cipher Attempt : {algo}"

        elif algo_type == "mac":
            title = f"SSH Weak MAC Attempt : {algo}"

        elif algo_type == "kex":
            title = f"SSH Weak Key Exchange Attempt : {algo}"

        elif algo_type == "host_key":
            title = f"SSH Weak Host Key Attempt : {algo}"

        else:
            title = f"SSH Weak Algorithm Attempt : {algo}"

        # -------- Overview --------
        if algo_type == "cipher":

            overview = (
                f"The above screenshot shows the SSH negotiation attempt where the weak "
                f"encryption algorithm {algo} was forced during the SSH communication between "
                f"the tester system and the DUT. The result shows that the algorithm was {negotiation_text}. "
                f"According to ITSAR guidelines, this algorithm is classified as {status_label}."
            )

        elif algo_type == "mac":

            overview = (
                f"The above screenshot shows the SSH negotiation attempt where the weak MAC algorithm {algo} "
                f"was forced. The negotiation result indicates that the algorithm was {negotiation_text}. "
                f"According to ITSAR guidelines, this MAC is classified as {status_label}."
            )

        elif algo_type == "kex":

            overview = (
                f"The above screenshot shows the SSH handshake where the weak key exchange algorithm {algo} "
                f"was attempted. The negotiation result shows that the algorithm was {negotiation_text}. "
                f"According to ITSAR guidelines, this KEX is classified as {status_label}."
            )

        elif algo_type == "host_key":

            overview = (
                f"The above screenshot shows the SSH handshake where the weak host key algorithm {algo} "
                f"was attempted. The negotiation result indicates that the algorithm was {negotiation_text}. "
                f"According to ITSAR guidelines, this host key is classified as {status_label}."
            )

        else:

            overview = (
                f"The above screenshot shows the SSH negotiation attempt using the algorithm {algo}. "
                f"The result indicates that the algorithm was {negotiation_text}. "
                f"According to ITSAR guidelines, this algorithm is classified as {status_label}."
            )

        # -------- Screenshot Block (UPDATED) --------
        if os.path.exists(img):

            add_screenshot_evidence_block(
                doc,
                title,
                img,
                overview   # 🔥 IMPORTANT (inside box)
            )

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(18)

    add_bold_paragraph(doc, "d) Test Observations:")

    results = weak_cipher_result.get("results", [])

# Find negotiated weak algorithms
    negotiated_algos = [
    f"{res.get('type')}:{res.get('algorithm')}"
    for res in results if res.get("negotiated")
    ]

    if negotiated_algos:
     negotiated_list = ", ".join(negotiated_algos)

     observation = (
        f"The SSH negotiation tests indicate that weak or deprecated cryptographic algorithm(s) "
        f"{negotiated_list} were successfully negotiated during the SSH handshake process. "
        f"As per ITSAR cryptographic requirements, such algorithms must be strictly rejected. "
        f"The acceptance of these algorithms indicates non-compliance with security requirements. "
        f"Hence, the DUT is marked as FAIL for this test case."
     )
    else:
     observation = (
        "The SSH negotiation tests indicate that all attempted weak or deprecated cryptographic algorithms "
        "were rejected by the DUT during the SSH handshake process. "
        "No weak algorithm was successfully negotiated. "
        "This confirms that the DUT enforces secure cryptographic policies as per ITSAR requirements. "
        "Hence, the DUT is marked as PASS for this test case."
     )

    doc.add_paragraph(observation)

    add_bold_paragraph(doc, "d) Evidence Provided :")
    doc.add_paragraph("Screenshots and Packet captures are Provided during Testing.")

# ================= NONE CIPHER ATTEMPT =================

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 4", 2)
    section_counter += 1
    secondsection_counter = 1
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC4_SSH_NO_ENCRYPTION_REJECTION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph("Verify that SSHv2 communication on the DUT fails when attempting to establish a session without encryption (e.g., none cipher)."
    "The DUT shall enforce encryption and must not allow unencrypted SSH communication.")

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph(
    "• The tester shall attempt to initiate an SSH connection to the DUT by explicitly forcing the use of no encryption (e.g., 'none' cipher)."
    )
    doc.add_paragraph(
    "• The tester shall observe the SSH client output and Server output to verify whether the connection attempt is rejected due to unsupported or disallowed cipher configuration."
    )
    doc.add_paragraph(
    "• Additionally, the tester shall verify through SSH algorithm enumeration that the DUT does not advertise or support the 'none' cipher in its encryption algorithms."
    )


    add_bold_paragraph(doc, "Execution Command:")
    doc.add_paragraph(none_cipher_result["user_input"])

    add_bold_paragraph(doc, "Executed Command Output:")
    doc.add_paragraph(none_cipher_result["terminal_output"] or "No output")
 
    if none_cipher_result.get("screenshot") and os.path.exists(none_cipher_result["screenshot"]):

     terminal_output = none_cipher_result.get("terminal_output", "").strip()
     remarks = none_cipher_result.get("remarks", "")
     none_supported = none_cipher_result.get("None_cipher_exist", False)

    # -------- Overview --------
     overview = (
        f"The above screenshot shows the SSH connection attempt where the tester explicitly forced the use of "
        f"no encryption (i.e., 'none' cipher) while initiating communication with the DUT. The output indicates: {remarks}. "
     )

    if none_supported:
        overview += (
            "The presence of 'none' cipher in the supported encryption list indicates that the DUT allows unencrypted communication, "
            "which is a critical security vulnerability and violates ITSAR cryptographic requirements."
        )

    # -------- Screenshot Block (UPDATED) --------
    add_screenshot_evidence_block(
        doc,
        "Test Case 4 : SSH None Cipher Attempt",
        none_cipher_result["screenshot"],
        overview   # 🔥 INSIDE BOX
    )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


    add_bold_paragraph(doc, "d) Test Observations:")

    result = none_cipher_result.get("result", "UNKNOWN")
    remarks = none_cipher_result.get("remarks", "")
    none_supported = none_cipher_result.get("None_cipher_exist", False)

    if result == "PASS":
     observation = (
        f"The SSH connection attempt using the 'none' cipher was not successful. {remarks}. "
        f"This ensures that unencrypted communication is not permitted by the DUT. "
        f"Hence, the DUT complies with ITSAR cryptographic requirements and is marked as PASS for this test case."
     )
    else:
     observation = (
        f"The SSH connection attempt using the 'none' cipher indicates that {remarks}. "
        f"This suggests that unencrypted communication may be possible, which violates ITSAR cryptographic requirements. "
        f"Hence, the DUT is marked as FAIL for this test case."
     )

    doc.add_paragraph(observation)

    add_bold_paragraph(doc, "d) Evidence Provided :")
    doc.add_paragraph("Screenshots and Packet captures are Provided during Testing.")




def add_https_section(doc, https_cipher_data, https_data, https_weak_cipher_result, tc8_result, section_counter,secondsection_counter,thirdsection_counter):

    add_itsar_heading(doc, f"{section_counter}. Test Execution For HTTPS", 2)

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 5", 2)

    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph("DUT should support only strong TLS cryptographic ciphers for HTTPS communication")

    add_bold_paragraph(doc, "c) Execution Steps:")

    p = doc.add_paragraph()

    p.add_run("• The tester should run the command ")
    bold_run = p.add_run("nmap --script ssl-enum-ciphers -p 443 <ip address>")
    bold_run.bold = True
    p.add_run(" to enumerate the TLS ciphers supported by the DUT.")

    doc.add_paragraph("• Validate that only strong TLS ciphers are supported.")
    doc.add_paragraph("• Capture evidence screenshots.")

    add_bold_paragraph(doc, "Execution Command:")
    doc.add_paragraph(https_cipher_data.get("user_input", "Not Available"))

    add_bold_paragraph(doc, "Executed Command Output:")
    doc.add_paragraph(https_cipher_data.get("terminal_output", "No output"))

    details = https_cipher_data.get("details", {})

# ================= TLSv1.2 =================

    tls12 = details.get("TLSv1.2", {})

    h = add_itsar_subheading(
     doc,
     f"{section_counter}.{secondsection_counter}.{thirdsection_counter}. DUT-Supported TLSv1.2 Cipher Suites",
     2
    )
    thirdsection_counter += 1
    keep_with_next(h)

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    t.cell(0,0).text = "Strong Cipher Suites"
    t.cell(0,1).text = "Weak Cipher Suites"

    row = t.add_row().cells
    row[0].text = "\n".join(
    normalize_list(tls12.get("ciphers", {}).get("strong", []))
    )  
    row[1].text = "\n".join(
    normalize_list(tls12.get("ciphers", {}).get("weak", []))
    )

    prevent_table_row_split(t)


# ================= TLSv1.3 =================

    tls13 = details.get("TLSv1.3", {})

    h = add_itsar_subheading(
     doc,
     f"{section_counter}.{secondsection_counter}.{thirdsection_counter}. DUT-Supported TLSv1.3 Cipher Suites",
     2
    )
    thirdsection_counter = 1
    secondsection_counter += 1
    keep_with_next(h)

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    t.cell(0,0).text = "Strong Cipher Suites"
    t.cell(0,1).text = "Weak Cipher Suites"

    row = t.add_row().cells
    row[0].text = "\n".join(
    normalize_list(tls13.get("ciphers", {}).get("strong", []))
    )
    row[1].text = "\n".join(
    normalize_list(tls13.get("ciphers", {}).get("weak", []))
    )

    prevent_table_row_split(t)

    doc.add_paragraph()

# ================= SCREENSHOT =================

    if https_cipher_data.get("screenshot") and os.path.exists(https_cipher_data["screenshot"]):

        add_screenshot_evidence_block(
            doc,
            "Test Case 1 : Executed cmd output: nmap --script ssl-enum-ciphers -p 443",
            https_cipher_data["screenshot"]
        )

# ================= HTTPS TC-02 =================

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter}. Test Case Number: 6", 2)
    secondsection_counter += 1
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC6_HTTPS_SECURE_COMMUNICATION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph(
        "DUT and tester devices should support traffic protection through HTTPS "
        "using secure TLS cryptographic protocols."
    )

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph("• The tester should initiate HTTPS communication with the DUT using openssl.")
    doc.add_paragraph("• Capture the TLS handshake packets using t-shark.")
    doc.add_paragraph(
        "• Analyze the Server Hello packet in the captured pcap file to identify "
        "the TLS protocol version and cipher used during communication."
    )

# ---------------- EXECUTION COMMAND ----------------

    add_bold_paragraph(doc, "Execution Command:")
    doc.add_paragraph(https_data.get("user_input", "Not Available"))

    add_bold_paragraph(doc, "Executed Command Output:")
    doc.add_paragraph(https_data.get("terminal_output", "No output"))

# ---------------- TLS CRYPTO TABLE ----------------

    crypto = https_data.get("crypto_details", {})

    h = add_itsar_subheading(doc,f"{section_counter}.{secondsection_counter}.{thirdsection_counter}. Security Ciphers Utilized During HTTPS Data Transmission",2)
    keep_with_next(h)

    tls_table = doc.add_table(rows=2, cols=2)
    tls_table.style = "Table Grid"

    tls_table.cell(0,0).text = "Protocol"
    tls_table.cell(0,1).text = crypto.get("protocol","Not Found")

    tls_table.cell(1,0).text = "Encryption Algorithm"
    tls_table.cell(1,1).text = crypto.get("cipher","Not Found")

    prevent_table_row_split(tls_table)

    p = doc.add_paragraph()
    keep_with_next(p)

# ---------------- SCREENSHOTS ----------------

    protocol = crypto.get("protocol", "Unknown")
    cipher = crypto.get("cipher", "Unknown")

    proto_status = https_data.get("nist_validation", {}).get("protocol", "FAIL")
    cipher_status = https_data.get("nist_validation", {}).get("cipher", "FAIL")

    proto_label = "Secure" if proto_status == "PASS" else "Insecure"
    cipher_label = "Secure" if cipher_status == "PASS" else "Insecure"

    if https_data.get("screenshots"):

     titles = [
        "Test Case 2 : HTTPS CLI Packet Capture showing TLS Handshake Traffic",
        f"Test Case 2 : TLS Server Hello showing {cipher_label} Cipher : {cipher}",
     ]

     overviews = [
        "The above screenshot shows the command line packet capture using t-shark while performing HTTPS communication between the tester system and the DUT. The captured packets include the TLS handshake exchanged during session establishment.",
        
        f"The above screenshot shows the TLS Server Hello packet captured during the HTTPS handshake process. "
        f"The Server Hello confirms that TLS version {protocol} and cipher suite {cipher} are used to secure the communication. "
        f"According to ITSAR cryptographic guidelines, this configuration is classified as {cipher_label}."
     ]

    for idx, (img, title, overview) in enumerate(zip(https_data["screenshots"], titles, overviews)):

        if os.path.exists(img):

            add_screenshot_evidence_block(
                doc,
                title,
                img,
                overview   # 🔥 INSIDE BOX
            )

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(18)
    


    # ================= HTTPS WEAK CIPHER REJECTION =================

    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 7", 2)
    secondsection_counter += 1
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC7_HTTPS_WEAK_CIPHER_REJECTION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph(
    "Verify that HTTPS communication over TLS on the DUT fails when attempting to establish a session "
    "using weak or deprecated cipher suites. The DUT shall reject negotiation of such weak cipher suites."
    )

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph(
    "• The tester shall initiate HTTPS connection attempts to the DUT using weak or deprecated TLS cipher suites."
    )
    doc.add_paragraph(
    "• OpenSSL shall be used to force specific weak cipher suites during the TLS handshake process."
    )
    doc.add_paragraph(
    "• The handshake output shall be analyzed to determine whether the weak cipher suite is accepted or rejected by the DUT."
    )

# ================= RESULT PROCESSING =================

    results = https_weak_cipher_result.get("results", [])
    screenshots = https_weak_cipher_result.get("screenshots", [])

    for res, img in zip(results, screenshots):

     cmd = res.get("command", "N/A")
     terminal_output = res.get("terminal_output", "").strip()
     cipher = res.get("cipher", "Unknown")
     tls_version = res.get("tls_version", "Unknown")
     negotiated = res.get("negotiated", False)

    # -------- Execution Command --------
     add_bold_paragraph(doc, f"Execution Command for weak TLS cipher: {cipher}")
     doc.add_paragraph(cmd)

    # -------- Command Output --------
     add_bold_paragraph(doc, "Executed Command Output:")

     if terminal_output:
        doc.add_paragraph(terminal_output)
     else:
        doc.add_paragraph("No relevant output captured (cipher likely rejected during negotiation)")

    # -------- Title --------
     title = f"HTTPS Weak Cipher Attempt ({tls_version}) : {cipher}"

     negotiation_text = "successfully negotiated" if negotiated else "rejected by the DUT"
     status_label = "Insecure"

    # -------- Overview --------
     overview = (
        f"The above screenshot shows the TLS handshake attempt where the weak cipher suite {cipher} "
        f"was forced during HTTPS communication between the tester system and the DUT using {tls_version}. "
        f"The result of the negotiation shows that the cipher suite was {negotiation_text}. "
        f"According to the cryptographic ITSAR guidelines, the cipher suite {cipher} is classified as {status_label}."
     )

    # -------- Screenshot Block --------
     if img and os.path.exists(img):

        add_screenshot_evidence_block(
            doc,
            title,
            img,
            overview
        )

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)

# ================= OBSERVATION =================

    add_bold_paragraph(doc, "d) Test Observations:")

    negotiated_ciphers = [
     res.get("cipher")
     for res in results if res.get("negotiated")
    ]

    if negotiated_ciphers:

     negotiated_list = ", ".join(negotiated_ciphers)

     observation = (
        f"The HTTPS negotiation tests indicate that weak or deprecated TLS cipher suite(s) "
        f"{negotiated_list} were successfully negotiated during the TLS handshake process. "
        f"As per ITSAR cryptographic requirements, such cipher suites must be strictly rejected. "
        f"The acceptance of these cipher suites indicates non-compliance with security requirements. "
        f"Hence, the DUT is marked as FAIL for this test case."
     )

    else:

     observation = (
        "The HTTPS negotiation tests indicate that all attempted weak or deprecated TLS cipher suites "
        "were rejected by the DUT during the TLS handshake process. "
        "No weak cipher suite was successfully negotiated. "
        "This confirms that the DUT enforces secure cryptographic policies as per ITSAR requirements. "
        "Hence, the DUT is marked as PASS for this test case."
     )

    doc.add_paragraph(observation)

    add_bold_paragraph(doc, "e) Evidence Provided:")
    doc.add_paragraph("Screenshots and command outputs are provided during testing.")


    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 8", 2)
    section_counter += 1
    secondsection_counter = 1
    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC8_HTTPS_NO_ENCRYPTION_REJECTION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph(
    "Verify that HTTPS communication on the DUT fails when attempting to establish a session using NULL or non-encrypted cipher suites. "
    "The DUT shall enforce encryption and must not allow unencrypted HTTPS communication across all supported TLS versions."
    )

# ================= EXECUTION STEPS =================

    add_bold_paragraph(doc, "c) Execution Steps:")

    doc.add_paragraph(
    "• The tester shall attempt to initiate an HTTPS connection to the DUT by explicitly forcing the use of NULL cipher suites using TLS 1.2."
    )

    doc.add_paragraph(
    "• The tester shall observe the OpenSSL output to verify whether the TLS handshake is rejected due to unsupported or disallowed cipher configuration."
    )

    doc.add_paragraph(
    "• The tester shall also initiate an HTTPS connection using TLS 1.3 to verify that only strong encryption cipher suites are negotiated, as TLS 1.3 does not support NULL ciphers."
    )

# ================= TLS 1.2 =================

    tls12_remarks = tc8_result["tls1_2"].get("remarks", "")

    overview_tls12 = (
    f"The above screenshot shows the HTTPS connection attempt where the tester forced the use of NULL cipher suites "
    f"using TLS 1.2. The result indicates: {tls12_remarks}. "
    f"This confirms that the DUT rejects unencrypted communication attempts and enforces secure cryptographic controls."
    )

    add_bold_paragraph(doc, "Execution Command (TLS 1.2):")
    doc.add_paragraph(tc8_result["tls1_2"]["command"])

    add_bold_paragraph(doc, "Executed Command Output (TLS 1.2):")
    doc.add_paragraph(tc8_result["tls1_2"]["output"] or "No output")

    if tc8_result["tls1_2"].get("screenshot") and os.path.exists(tc8_result["tls1_2"]["screenshot"]):

     add_screenshot_evidence_block(
        doc,
        "Test Case 8 : HTTPS NULL Cipher Attempt (TLS 1.2)",
        tc8_result["tls1_2"]["screenshot"],
        overview_tls12   # 🔥 INSIDE BOX
     )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


# ================= TLS 1.3 =================

    tls13_remarks = tc8_result["tls1_3"].get("remarks", "")

    overview_tls13 = (
    f"The above screenshot shows the HTTPS connection established using TLS 1.3. "
    f"The negotiation confirms that only strong cipher suites are enforced, indicating: {tls13_remarks}. "
    f"As TLS 1.3 does not support NULL or weak ciphers, secure encryption is inherently enforced."
    )

    add_bold_paragraph(doc, "Execution Command (TLS 1.3):")
    doc.add_paragraph(tc8_result["tls1_3"]["command"])

    add_bold_paragraph(doc, "Executed Command Output (TLS 1.3):")
    doc.add_paragraph(tc8_result["tls1_3"]["output"] or "No output")

    if tc8_result["tls1_3"].get("screenshot") and os.path.exists(tc8_result["tls1_3"]["screenshot"]):

     add_screenshot_evidence_block(
        doc,
        "Test Case 8 : HTTPS Secure Connection (TLS 1.3)",
        tc8_result["tls1_3"]["screenshot"],
        overview_tls13   # 🔥 INSIDE BOX
     )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

# ================= OBSERVATIONS =================

    add_bold_paragraph(doc, "d) Test Observations:")

    tls12_result = tc8_result["tls1_2"].get("result", "UNKNOWN")
    tls13_result = tc8_result["tls1_3"].get("result", "UNKNOWN")
    final_result = tc8_result.get("final_result", "UNKNOWN")

    if final_result == "PASS":
     observation = (
        "The HTTPS connection attempt using NULL cipher suites under TLS 1.2 was not successful, confirming that the DUT rejects unencrypted communication. "
        "Additionally, under TLS 1.3, the DUT successfully enforces strong cipher suites and does not allow insecure configurations. "
        "This ensures that all HTTPS communication is encrypted and secure. "
        "Hence, the DUT complies with ITSAR cryptographic requirements and is marked as PASS for this test case."
     )
    else:
     observation = (
        "The HTTPS connection attempt using NULL or non-encrypted cipher suites indicates that unencrypted or insecure communication may be possible. "
        "This violates ITSAR cryptographic requirements, as secure communication must be enforced at all times. "
        "Hence, the DUT is marked as FAIL for this test case."
     )

    doc.add_paragraph(observation)

# ================= EVIDENCE =================

    add_bold_paragraph(doc, "e) Evidence Provided :")
    doc.add_paragraph("Screenshots and terminal outputs are provided during testing.")

    


def add_snmp_section(doc,snmp_tc1_data,snmp_tc2_data,section_counter,secondsection_counter):

    add_itsar_heading(doc, f"{section_counter}. Test Execution For SNMP", 2)
   
    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 9", 2)
    secondsection_counter += 1

    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC9_SNMP_VERSION_CHECK")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph(
        "Verify that the DUT does not allow insecure SNMP versions (SNMPv1 and SNMPv2c). "
        "The DUT shall only support SNMPv3 for secure communication. "
        "Any successful response using SNMPv1 or SNMPv2c indicates insecure configuration."
        )

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph("• Attempt SNMPv1 communication using default community string.")
    doc.add_paragraph("• Attempt SNMPv2c communication using default community string.")
    doc.add_paragraph("• Observe whether the DUT responds with valid SNMP data.")
    doc.add_paragraph("• Successful response indicates insecure configuration.")

# ---------------- SNMPv1 ----------------

    add_bold_paragraph(doc, "Execution Command (SNMPv1):")
    doc.add_paragraph(snmp_tc1_data["user_input_v1"])

    add_bold_paragraph(doc, "Executed Command Output (SNMPv1):")
    doc.add_paragraph(snmp_tc1_data["terminal_output_v1"] or "No output")

    v1_success = snmp_tc1_data.get("validation_details", {}).get("v1_success", False)

    if snmp_tc1_data.get("v1_screenshot") and os.path.exists(snmp_tc1_data["v1_screenshot"]):

        if v1_success:
            overview_v1 = (
                "The above screenshot shows the SNMPv1 communication attempt with the DUT. "
                "The output indicates that the DUT responded with valid SNMP data using SNMPv1, "
                "which confirms that SNMPv1 is enabled. This is considered insecure as SNMPv1 uses "
                "plaintext community strings and does not provide authentication or encryption. "
                "Hence, SNMPv1 is improperly enabled on the DUT."
            )
        else:
            overview_v1 = (
                "The above screenshot shows the SNMPv1 communication attempt with the DUT. "
                "The output indicates that no response was received from the DUT, confirming that SNMPv1 is disabled. "
                "This ensures that insecure SNMPv1 communication is not permitted, which aligns with ITSAR security requirements."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 9 : SNMPv1 Communication Attempt",
            snmp_tc1_data["v1_screenshot"],
            overview_v1
        )

    # ---------------- SNMPv2c ----------------

    add_bold_paragraph(doc, "Execution Command (SNMPv2c):")
    doc.add_paragraph(snmp_tc1_data["user_input_v2c"])

    add_bold_paragraph(doc, "Executed Command Output (SNMPv2c):")
    doc.add_paragraph(snmp_tc1_data["terminal_output_v2c"] or "No output")

    v2_success = snmp_tc1_data.get("validation_details", {}).get("v2c_success", False)

    if snmp_tc1_data.get("v2c_screenshot") and os.path.exists(snmp_tc1_data["v2c_screenshot"]):

        if v2_success:
            overview_v2 = (
                "The above screenshot shows the SNMPv2c communication attempt with the DUT. "
                "The output indicates that the DUT responded with valid SNMP data using SNMPv2c, "
                "which confirms that SNMPv2c is enabled. SNMPv2c relies on community strings and does not provide encryption, "
                "making it vulnerable to interception and unauthorized access. "
                "Hence, SNMPv2c is improperly enabled on the DUT."
            )
        else:
            overview_v2 = (
                "The above screenshot shows the SNMPv2c communication attempt with the DUT. "
                "The output indicates that no response was received from the DUT, confirming that SNMPv2c is disabled. "
                "This ensures that insecure SNMPv2c communication is not permitted, which aligns with ITSAR security requirements."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 9 : SNMPv2c Communication Attempt",
            snmp_tc1_data["v2c_screenshot"],
            overview_v2
        )

    # ---------------- OBSERVATION ----------------

    add_bold_paragraph(doc, "d) Test Observations:")

    result = snmp_tc1_data.get("final_result", "UNKNOWN")

    if result == "PASS":
        observation = (
            "The SNMPv1 and SNMPv2c communication attempts were not successful. "
            "This confirms that insecure SNMP versions are disabled and only secure SNMPv3 communication is allowed. "
            "Hence, the DUT complies with ITSAR cryptographic requirements and is marked as PASS for this test case."
        )
    else:
        observation = (
            "The SNMPv1 and/or SNMPv2c communication attempts were successful, indicating that the DUT allows insecure SNMP protocols. "
            "This exposes management information over unencrypted channels and violates ITSAR security requirements. "
            "Hence, the DUT is marked as FAIL for this test case."
        )

    doc.add_paragraph(observation)

    # ---------------- EVIDENCE ----------------

    add_bold_paragraph(doc, "e) Evidence Provided:")
    doc.add_paragraph("Screenshots of SNMPv1 and SNMPv2c communication attempts are provided as evidence.")


# -------------------------- SNMP TC2---------------------------------
    add_itsar_subheading(doc, f"{section_counter}.{secondsection_counter} Test Case Number: 10", 2)
    secondsection_counter = 1

    add_bold_paragraph(doc, "a) Test Case Name:")
    doc.add_paragraph("TC10_SNMP_SECURE_COMMUNICATION")

    add_bold_paragraph(doc, "b) Test Case Description:")
    doc.add_paragraph(
        "Verify that SNMPv3 communication on the DUT enforces secure authentication and encryption (authPriv mode). "
        "The DUT shall allow only authPriv mode and must reject weaker modes such as authNoPriv and noAuthNoPriv."
    )

    add_bold_paragraph(doc, "c) Execution Steps:")
    doc.add_paragraph("• Initiate SNMPv3 communication using authPriv (SHA + AES).")
    doc.add_paragraph("• Capture packets in Wireshark and verify encryption and authentication flags.")
    doc.add_paragraph("• Attempt SNMPv3 communication using authNoPriv.")
    doc.add_paragraph("• Attempt SNMPv3 communication using noAuthNoPriv.")
    doc.add_paragraph("• Verify that only authPriv mode is allowed.")

    # ================= AUTHPRIV =================

    authPriv = snmp_tc2_data["authPriv"]

    add_bold_paragraph(doc, "Execution Command (authPriv):")
    doc.add_paragraph(authPriv["command"])

    add_bold_paragraph(doc, "Executed Command Output (authPriv):")
    doc.add_paragraph(authPriv["output"] or "No output")

    if authPriv.get("terminal_screenshot") and os.path.exists(authPriv["terminal_screenshot"]):

        if authPriv["success"]:
            overview = (
                "The above screenshot shows the SNMPv3 communication attempt using authPriv mode. "
                "The DUT responded successfully, confirming that secure SNMP communication is supported."
            )
        else:
            overview = (
                "The above screenshot shows the SNMPv3 communication attempt using authPriv mode. "
                "The DUT did not respond successfully, indicating improper secure SNMP configuration."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 authPriv Terminal Output",
            authPriv["terminal_screenshot"],
            overview
        )

    if authPriv.get("wireshark_screenshot") and os.path.exists(authPriv["wireshark_screenshot"]):

        if authPriv["success"]:
            ws_overview = (
                "The above Wireshark capture shows SNMPv3 packet details. "
                "The msgFlags field indicates that both authentication and encryption are enabled, "
                "confirming that authPriv mode is correctly enforced using secure mechanisms (SHA and AES)."
            )
        else:
            ws_overview = (
                "The above Wireshark capture shows SNMPv3 packet details. "
                "The expected authentication and encryption flags are not properly set, indicating that "
                "secure SNMPv3 communication (authPriv) is not correctly enforced on the DUT."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 authPriv Wireshark Evidence",
            authPriv["wireshark_screenshot"],
            ws_overview
        )

    # ================= AUTHNOPRIV =================

    authNoPriv = snmp_tc2_data["authNoPriv"]

    add_bold_paragraph(doc, "Execution Command (authNoPriv):")
    doc.add_paragraph(authNoPriv["command"])

    add_bold_paragraph(doc, "Executed Command Output (authNoPriv):")
    doc.add_paragraph(authNoPriv["output"] or "No output")

    if authNoPriv.get("terminal_screenshot") and os.path.exists(authNoPriv["terminal_screenshot"]):

        if authNoPriv["success"]:
            overview = (
                "The above screenshot shows the SNMPv3 communication attempt using authNoPriv mode. "
                "The DUT responded successfully, indicating that SNMP communication without encryption is allowed. "
                "This is insecure as data is transmitted without confidentiality protection."
            )
        else:
            overview = (
                "The above screenshot shows the SNMPv3 communication attempt using authNoPriv mode. "
                "The DUT did not respond, confirming that insecure communication without encryption is restricted."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 authNoPriv Terminal Output",
            authNoPriv["terminal_screenshot"],
            overview
        )

    if authNoPriv.get("wireshark_screenshot") and os.path.exists(authNoPriv["wireshark_screenshot"]):

        if authNoPriv["success"]:
            ws_overview = (
                "The Wireshark capture shows that authentication is enabled but encryption is not set in the msgFlags field. "
                "This confirms that SNMP communication is occurring without encryption, which is insecure."
            )
        else:
            ws_overview = (
                "The Wireshark capture indicates that the SNMP communication attempt did not succeed, "
                "confirming that authNoPriv mode is not permitted by the DUT."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 authNoPriv Wireshark Evidence",
            authNoPriv["wireshark_screenshot"],
            ws_overview
        )

    # ================= NOAUTHNOPRIV =================

    noAuth = snmp_tc2_data["noAuthNoPriv"]

    add_bold_paragraph(doc, "Execution Command (noAuthNoPriv):")
    doc.add_paragraph(noAuth["command"])

    add_bold_paragraph(doc, "Executed Command Output (noAuthNoPriv):")
    doc.add_paragraph(noAuth["output"] or "No output")

    if noAuth.get("terminal_screenshot") and os.path.exists(noAuth["terminal_screenshot"]):

        if noAuth["success"]:
            overview = (
                "The above screenshot shows SNMPv3 communication without authentication and encryption. "
                "The DUT responded successfully, indicating a critical security vulnerability."
            )
        else:
            overview = (
                "The above screenshot shows SNMPv3 communication attempt without authentication and encryption. "
                "The DUT did not respond, confirming that this insecure mode is properly disabled."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 noAuthNoPriv Terminal Output",
            noAuth["terminal_screenshot"],
            overview
        )

    if noAuth.get("wireshark_screenshot") and os.path.exists(noAuth["wireshark_screenshot"]):

        if noAuth["success"]:
            ws_overview = (
                "The Wireshark capture confirms that neither authentication nor encryption flags are set, "
                "indicating a highly insecure communication mode allowed by the DUT."
            )
        else:
            ws_overview = (
                "The Wireshark capture confirms that the SNMP communication attempt without authentication "
                "and encryption was not successful, indicating proper restriction of insecure modes."
            )

        add_screenshot_evidence_block(
            doc,
            "Test Case 10 : SNMPv3 noAuthNoPriv Wireshark Evidence",
            noAuth["wireshark_screenshot"],
            ws_overview
        )

    # ================= OBSERVATION =================

    add_bold_paragraph(doc, "d) Test Observations:")

    result = snmp_tc2_data.get("final_result", "UNKNOWN")

    if result == "PASS":
        observation = (
            "The DUT allows only secure SNMPv3 communication using authPriv mode and rejects weaker modes. "
            "This ensures confidentiality and integrity of SNMP communication. "
            "Hence, the DUT complies with ITSAR requirements and is marked as PASS."
        )
    else:
        observation = (
            "The DUT allows SNMP communication using weaker security modes such as authNoPriv, "
            "which does not provide encryption. This exposes sensitive data over the network and "
            "violates ITSAR security requirements. Hence, the DUT is marked as FAIL."
        )

    doc.add_paragraph(observation)

    # ================= EVIDENCE =================

    add_bold_paragraph(doc, "e) Evidence Provided:")
    doc.add_paragraph("Terminal outputs and Wireshark captures are provided as evidence for all SNMPv3 modes tested.")
   

if __name__ == "__main__":
        

    test_Data = {
        "details": {
            "encryption": {"weak":[]},
            "mac": {"weak":[]},
            "kex": {"weak":["ecdh-sha2-nistp521"]},
            "host_key": {"weak":[]}
        },
    }

    test_data = {
       "details": {
             "TLSv1.2": {
               "ciphers": {
                  "weak": [
                    "ECDHE-RSA-AES128-SHA",
                  ]
             }
         },
         "TLSv1.3": {
            "ciphers": {
                "weak": [
                    "TLS_AES_128_GCM_SHA256",
                    "TLS_AES_128_CCM_SHA256"
                ]
            }
        }
          }
           
    }

    # -------- Run Nmap Scan --------
    nmap_data = run_nmap_scan()


    # # # # -------- OEM Verification ------
    OEM_data = run_OEM_test()

    # # # # -------- Initialize Empty Results --------
    cipher_result = {}
    ssh_result = {}
    weak_cipher_result = {}
    none_cipher_result = {}

    https_cipher_data = {}
    https_data = {}
    https_weak_cipher_result = {}
    tc8_result = {}

    ssh_applicable = False
    https_applicable = False
    snmp_applicable = False


    # -------- SSH Service Detection --------
    if nmap_data.get("SSH"):

        ssh_applicable = True

        print("\n[+] SSH service detected — running SSH tests...\n")

        cipher_result = run_cipher_detection()
        ssh_result = run_ssh_verification()
        weak_cipher_result = run_ssh_weak_cipher_test(cipher_result,ssh_result)
        none_cipher_result = run_ssh_none_cipher_test(cipher_result)

    else:
        print("\n[-] SSH service NOT detected — skipping SSH tests\n")


    # # -------- HTTPS Service Detection --------
    if nmap_data.get("HTTPS"):

        https_applicable = True

        print("\n[+] HTTPS service detected — running HTTPS tests...\n")

        https_cipher_data = run_httpsCipher_detection()
        https_data = run_tls_verification()
        https_weak_cipher_result = run_https_weak_cipher_test(test_data)
        tc8_result = run_https_NULL_test()

    else:
        print("\n[-] HTTPS service NOT detected — skipping HTTPS tests\n")

    ## --------- SNMP Service Detection -----------
    if nmap_data.get("SNMP"):

        snmp_applicable = True

        print("\n[+] SNMP service detected — running SNMP tests...\n")

        snmp_tc1_data = run_snmp_tc1()
        snmp_tc2_data = run_snmp_tc2()
       

    else:
        print("\n[-] SNMP service NOT detected — skipping SNMP tests\n")


    # # # # -------- DUT Information --------
    dut_info = get_dut_info(
        os.getenv("DUT_USER"),
        os.getenv("DUT_IP")
    )

    # # -------- Generate Report --------
    report = create_word_report(
    nmap_data,
    OEM_data,
    cipher_result,
    ssh_result,
    weak_cipher_result,
    none_cipher_result,
    https_cipher_data,
    https_data,
    https_weak_cipher_result,
    tc8_result,
    dut_info,
    snmp_tc1_data,
    snmp_tc2_data,
    ssh_applicable,
    https_applicable,
    snmp_applicable
    )

    print(f"\n✅ REPORT GENERATED SUCCESSFULLY: {report}\n")